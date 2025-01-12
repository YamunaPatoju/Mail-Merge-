from docx import Document

PLACEHOLDER = "[name]"

# Read names from the text file and strip newlines
with open("C:\\Users\\DELL\\Desktop\\mail\\input\\names\\persons_names.txt") as names_file:
    names = [name.strip() for name in names_file.readlines()]

# Load the template letter document (corrected file extension)
doc = Document("C:\\Users\\DELL\\Desktop\\mail\\input\\letter\\starting_ll.docx.docx")

# Extract the content of the letter template
letters_content = "\n".join([para.text for para in doc.paragraphs])

# Loop through each name and create a personalized letter for each
for name in names:
    # Replace the placeholder with the current name
    new_letter_content = letters_content.replace(PLACEHOLDER, name)

    # Create a new Document for the current letter
    new_letter = Document()

    # Add the modified content to the new Document
    for paragraph in new_letter_content.split("\n"):
        new_letter.add_paragraph(paragraph)

    # Save the new document to the output folder with the person's name in the filename
    output_path = (f"C:\\Users\\DELL\\Desktop\\mail\\output\\Ready_to_start\\letter_for_{name}.docx")
    new_letter.save(output_path)

    # Optional: Print a success message for each saved file
    print(f"Saved letter for {name}")
